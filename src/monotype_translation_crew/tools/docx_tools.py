import io
import json
import logging
import os
import re
import threading
import zipfile
from pathlib import Path

from crewai.tools import tool

logger = logging.getLogger(__name__)


def _normalise_docx_translation(text: str, lang: str) -> str:
    """Mirror of excel_tools._normalise_translation for the DOCX write path."""
    had_dots = '...' in text
    text = text.replace('...', '…')
    text = re.sub(r'\{\{\s+(\S+?)\s+\}\}', r'{{\1}}', text)
    text = text.replace('\u2019', "'").replace('\u2018', "'")
    # German uses „” (U+201E/U+201C) as native quotes — do not convert for de
    if lang != "de":
        text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.rstrip()
    if lang == "fr":
        text = re.sub(r' ([?!:;])', r' \1', text)
    logger.debug("_normalise_docx_translation: lang=%s ellipsis_fixed=%s", lang, had_dots)
    return text

_DEFAULT_REVIEWED_DOCX_PATH = os.path.join("outputs", "reviewed_docx_translations.json")
_job_docx_local = threading.local()


def _get_reviewed_docx_path() -> str:
    return getattr(_job_docx_local, "path", _DEFAULT_REVIEWED_DOCX_PATH)


def set_job_docx_translation_path(path: str) -> None:
    """Bind a job-scoped reviewed_docx_translations path to the current thread."""
    _job_docx_local.path = path


def extract_segments(docx_path: str) -> list[dict]:
    """Pure-Python helper (not a tool): extract translatable segments from a docx.

    Returns the same list that read_docx_for_translation would embed in its JSON,
    but as a plain Python list so callers can inspect count before running agents.
    """
    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError:
        return []

    if not os.path.exists(docx_path):
        return []

    doc = Document(docx_path)
    segments: list[dict] = []
    seg_id = 0

    for tag, block in _get_body_blocks(doc):
        if tag == "p":
            para = Paragraph(block, doc)
            text = _para_text(para)
            if text:
                segments.append({"segment_id": seg_id, "type": "paragraph", "text": text})
                seg_id += 1
        elif tag == "tbl":
            table = Table(block, doc)
            for row_idx, row in enumerate(table.rows):
                seen_texts: set[str] = set()
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text and text not in seen_texts:
                        seen_texts.add(text)
                        segments.append({
                            "segment_id": seg_id,
                            "type": "table_cell",
                            "text": text,
                            "table_row": row_idx,
                            "table_col": col_idx,
                        })
                        seg_id += 1

    return segments

LANG_LABELS = {
    "fr": "French", "de": "German",
    "pt_BR": "Portuguese (Brazil)", "ja": "Japanese", "es_ES": "Spanish (Spain)",
}


def _get_body_blocks(doc):
    """Yield (tag, element) for each top-level body block."""
    from docx.oxml.ns import qn
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        yield tag, child


def _para_text(para) -> str:
    return "".join(r.text for r in para.runs).strip()


def _replace_para_text(para, new_text: str) -> None:
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _replace_cell_text(cell, new_text: str) -> None:
    if not cell.paragraphs:
        return
    _replace_para_text(cell.paragraphs[0], new_text)
    for para in cell.paragraphs[1:]:
        for run in para.runs:
            run.text = ""


@tool
def read_docx_for_translation(docx_path: str) -> str:
    """
    Reads a Word document and extracts all translatable text segments
    (body paragraphs and table cells) as a JSON structure.

    Args:
        docx_path: Absolute or relative path to the .docx file.

    Returns:
        JSON string with keys: docx_path, segments (list), total_segments.
        Each segment has: segment_id (int), type (paragraph|table_cell), text (str).
    """
    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError:
        return json.dumps({"error": "python-docx not installed. Run: pip install python-docx"})

    if not os.path.exists(docx_path):
        return json.dumps({"error": f"File not found: {docx_path}"})

    try:
        doc = Document(docx_path)
        segments = []
        seg_id = 0

        for tag, block in _get_body_blocks(doc):
            if tag == "p":
                para = Paragraph(block, doc)
                text = _para_text(para)
                if text:
                    segments.append({"segment_id": seg_id, "type": "paragraph", "text": text})
                    seg_id += 1
            elif tag == "tbl":
                table = Table(block, doc)
                for row_idx, row in enumerate(table.rows):
                    seen_texts: set[str] = set()
                    for col_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if text and text not in seen_texts:
                            seen_texts.add(text)
                            segments.append({
                                "segment_id": seg_id,
                                "type": "table_cell",
                                "text": text,
                                "table_row": row_idx,
                                "table_col": col_idx,
                            })
                            seg_id += 1

        return json.dumps({
            "docx_path": docx_path,
            "segments": segments,
            "total_segments": len(segments),
        }, ensure_ascii=False, indent=2)

    except Exception as exc:
        return json.dumps({"error": f"Failed to read docx: {exc}"})


def write_translations_to_docx_impl(docx_path: str, json_path: str | None = None) -> dict:
    """Core implementation — reads reviewed_docx_translations.json, writes per-language
    .docx files and bundles them into a .zip.  Returns a plain dict (not JSON string)
    so it can be called directly from Python without going through an AI agent."""
    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    json_path = json_path or _get_reviewed_docx_path()
    if not os.path.isabs(json_path):
        json_path = os.path.join(os.getcwd(), json_path)
    if not os.path.exists(json_path):
        return {"error": f"Reviewed translations not found: {json_path}"}

    try:
        raw = Path(json_path).read_text(encoding="utf-8").strip()
        if raw.startswith("```"):
            raw = "\n".join(raw.splitlines()[1:])
        if raw.endswith("```"):
            raw = "\n".join(raw.splitlines()[:-1])
        translations: list[dict] = json.loads(raw)
    except Exception as exc:
        return {"error": f"Failed to parse reviewed_docx_translations.json: {exc}"}

    if not isinstance(translations, list) or not translations:
        return {"error": "reviewed_docx_translations.json must be a non-empty JSON array"}

    skip_keys = {"segment_id", "english", "reviewer_note"}
    target_langs = sorted({k for entry in translations for k in entry if k not in skip_keys})
    if not target_langs:
        return {"error": "No language keys found in reviewed translations JSON"}

    trans_by_id: dict[int, dict] = {t["segment_id"]: t for t in translations if "segment_id" in t}

    if not os.path.exists(docx_path):
        return {"error": f"Source docx not found: {docx_path}"}

    outputs_dir = Path(os.getcwd()) / "outputs"
    outputs_dir.mkdir(exist_ok=True)
    stem = Path(docx_path).stem

    lang_paths: dict[str, str] = {}
    segments_written: dict[str, int] = {}

    for lang in target_langs:
        doc = Document(docx_path)
        seg_id = 0
        written = 0

        for tag, block in _get_body_blocks(doc):
            if tag == "p":
                para = Paragraph(block, doc)
                if _para_text(para):
                    entry = trans_by_id.get(seg_id, {})
                    translated = entry.get(lang, "")
                    if translated:
                        _replace_para_text(para, _normalise_docx_translation(translated, lang))
                        written += 1
                    seg_id += 1
            elif tag == "tbl":
                table = Table(block, doc)
                for row in table.rows:
                    seen_texts: set[str] = set()
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in seen_texts:
                            seen_texts.add(text)
                            entry = trans_by_id.get(seg_id, {})
                            translated = entry.get(lang, "")
                            if translated:
                                _replace_cell_text(cell, _normalise_docx_translation(translated, lang))
                                written += 1
                            seg_id += 1

        out_path = str(outputs_dir / f"{stem}_{lang}.docx")
        doc.save(out_path)
        lang_paths[lang] = out_path
        segments_written[lang] = written

    zip_path = str(outputs_dir / f"{stem}_translated.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for lang, path in lang_paths.items():
            label = LANG_LABELS.get(lang, lang)
            zf.write(path, arcname=f"{stem}_{label}.docx")

    return {
        "success": True,
        "zip_path": zip_path,
        "languages_written": list(lang_paths.keys()),
        "segments_written": segments_written,
    }


@tool
def write_translations_to_docx(docx_path: str) -> str:
    """
    Reads reviewed_docx_translations.json and writes one translated .docx per language,
    then bundles them into a .zip file.

    Args:
        docx_path: Path to the original source .docx file.

    Returns:
        JSON string with keys: success, zip_path, languages_written, segments_written.
    """
    return json.dumps(write_translations_to_docx_impl(docx_path), ensure_ascii=False, indent=2)
