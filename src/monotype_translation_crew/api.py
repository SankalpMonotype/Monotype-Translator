"""FastAPI web application for Monotype Translation Crew.

Exposes a single-page web UI where users can:
  - Upload an .xlsx file
  - Watch a progress indicator while the CrewAI pipeline runs
  - Download the translated Excel output
  - Browse a review table of all translated strings

Known limitation: reviewed_translations.json is a shared file in outputs/.
Running two jobs concurrently will cause the review data to be overwritten.
For single-user / small-team use this is fine.
"""

import asyncio
import json
import math
import re
import uuid
from concurrent.futures import ThreadPoolExecutor
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

# ---------------------------------------------------------------------------
# File-format converters  (PDF / DOCX → translation-ready .xlsx)
# ---------------------------------------------------------------------------

_XLSX_HEADERS = ["English", "French", "German", "Portuguese (pt-BR)", "Japanese", "Spanish (es-ES)"]


def _write_strings_xlsx(strings: list[str], xlsx_path: str) -> None:
    """Write a flat list of strings into a 6-column translation Excel."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Translations"
    ws.append(_XLSX_HEADERS)
    for s in strings:
        if s and s.strip():
            ws.append([s.strip(), "", "", "", "", ""])
    wb.save(xlsx_path)


def _pdf_to_xlsx(pdf_path: str, xlsx_path: str) -> None:
    import pdfplumber
    strings: list[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            for table in (page.extract_tables() or []):
                for row in table:
                    for cell in row:
                        if cell and cell.strip():
                            strings.append(cell.strip())
            if not strings:                # fallback: raw text lines
                text = page.extract_text() or ""
                strings.extend(l.strip() for l in text.splitlines() if l.strip())
    if not strings:
        raise ValueError("No readable text found in the PDF.")
    _write_strings_xlsx(strings, xlsx_path)


def _docx_to_xlsx(docx_path: str, xlsx_path: str) -> None:
    from docx import Document
    doc = Document(docx_path)
    strings: list[str] = []
    seen: set[str] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    strings.append(t)
    if len(strings) < 2:        # no useful tables — fall back to paragraphs
        strings = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not strings:
        raise ValueError("No readable text found in the document.")
    _write_strings_xlsx(strings, xlsx_path)


def _convert_to_xlsx(src: Path, dst: Path) -> None:
    """Convert PDF or DOCX to a translation-ready .xlsx at *dst*."""
    ext = src.suffix.lower()
    if ext == ".pdf":
        _pdf_to_xlsx(str(src), str(dst))
    elif ext in (".docx", ".doc"):
        _docx_to_xlsx(str(src), str(dst))
    else:
        raise ValueError(f"Unsupported file type: {ext}")

# ---------------------------------------------------------------------------
# Job store (in-memory; lost on server restart)
# ---------------------------------------------------------------------------
JOBS: dict[str, dict] = {}

# Thread pool — CrewAI is synchronous, so we run it off the event loop
_executor = ThreadPoolExecutor(max_workers=2)


# ---------------------------------------------------------------------------
# App lifecycle
# ---------------------------------------------------------------------------

@asynccontextmanager
async def lifespan(app: FastAPI):
    Path("uploads").mkdir(exist_ok=True)
    Path("outputs").mkdir(exist_ok=True)
    yield


app = FastAPI(title="Monotype Translation Crew", version="1.0.0", lifespan=lifespan)
_STATIC_DIR = Path(__file__).parent / "static"
if _STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(_STATIC_DIR)), name="static")


# ---------------------------------------------------------------------------
# Embedded single-page frontend
# ---------------------------------------------------------------------------

_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Monotype Translation Crew</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;700;800&display=swap" rel="stylesheet">
  <style>
    @font-face { font-family:'Helvetica Now Display'; src:url('/static/fonts/HelveticaNowMTDisp.otf') format('opentype'); font-weight:400; font-style:normal; font-display:swap; }
    @font-face { font-family:'Helvetica Now Display'; src:url('/static/fonts/HelveticaNowMTDispMedium.otf') format('opentype'); font-weight:500; font-style:normal; font-display:swap; }
    @font-face { font-family:'Helvetica Now Display'; src:url('/static/fonts/HelveticaNowMTDispBold.otf') format('opentype'); font-weight:700; font-style:normal; font-display:swap; }
    @font-face { font-family:'Helvetica Now Text'; src:url('/static/fonts/HelveticaNowMTText.otf') format('opentype'); font-weight:400; font-style:normal; font-display:swap; }
    @font-face { font-family:'Helvetica Now Text'; src:url('/static/fonts/HelveticaNowMTTextMedium.otf') format('opentype'); font-weight:500; font-style:normal; font-display:swap; }
    @font-face { font-family:'Helvetica Now Text'; src:url('/static/fonts/HelveticaNowMTTextBold.otf') format('opentype'); font-weight:700; font-style:normal; font-display:swap; }

    :root {
      --font-display: 'Helvetica Now Display','HelveticaNowMTDisp',Inter,Helvetica,Arial,sans-serif;
      --font-text:    'Helvetica Now Text','HelveticaNowMTText',Inter,Helvetica,Arial,sans-serif;
      --font-mono:    ui-monospace,'Cascadia Code','Source Code Pro',monospace;

      /* Antiqua — Spirits neutrals */
      --s25:  #fcfcfd; --s50:  #f9fafb; --s100: #f3f4f6;
      --s200: #e7eaee; --s300: #dbdfe5; --s400: #cfd5dd;
      --s500: #c3cad4; --s600: #667488; --s700: #576579;
      --s800: #313945; --s900: #1e242c;

      /* Antiqua — Blue Duck primary */
      --bd-50:  #f2f5fd; --bd-100: #e4e9fa; --bd-200: #c6d2f6;
      --bd-500: #1a73e8; --bd-600: #1766cf; --bd-700: #1459b3;

      /* Antiqua — Status */
      --green-50:  #f2f7f4; --green-200: #c6dfd0;
      --green-500: #26a568; --green-700: #1d7f50;
      --red-50:    #fbf1f2; --red-200:   #f1c5c6; --red-500: #dc0024;

      --border: var(--s300);

      /* JS-inline-style compatibility aliases */
      --ink:      var(--s900);
      --paper:    var(--s50);
      --mist:     var(--bd-50);
      --muted:    var(--s700);
      --subtle:   var(--s600);
      --green:    var(--green-500);
      --green-bg: var(--green-50);
      --green-bd: var(--green-200);
    }

    *, *::before, *::after { box-sizing: border-box; }
    html, body { margin:0; padding:0; font-family:var(--font-text); color:var(--s900); background:var(--s100); -webkit-font-smoothing:antialiased; min-height:100vh; }

    /* ── Animations (functional only) ── */
    @keyframes spin      { to { transform:rotate(360deg); } }
    @keyframes pulse-dot { 0%,100%{opacity:1;} 50%{opacity:.35;} }
    @keyframes fade-up   { from{opacity:0;transform:translateY(6px);} to{opacity:1;transform:translateY(0);} }
    @keyframes draw-check{ to { stroke-dashoffset:0; } }

    .anim-fade-up  { animation:fade-up .18s cubic-bezier(0,0,.2,1) both; }
    .anim-slide-up { animation:fade-up .18s cubic-bezier(0,0,.2,1) both; }

    /* ── Step indicator ── */
    .step-indicator { display:flex; align-items:center; margin-bottom:28px; }
    .si-step  { display:flex; align-items:center; gap:6px; }
    .si-dot   { width:6px; height:6px; border-radius:50%; background:var(--s300); flex-shrink:0; transition:background .2s; }
    .si-step.si-active .si-dot { background:var(--s900); }
    .si-step.si-done   .si-dot { background:var(--s500); }
    .si-label { font-size:11px; color:var(--s600); white-space:nowrap; }
    .si-step.si-active .si-label { color:var(--s900); font-weight:500; }
    .si-step.si-done   .si-label { color:var(--s600); }
    .si-line  { flex:1; height:1px; background:var(--border); margin:0 8px; min-width:12px; max-width:32px; }

    /* ── Drop zone ── */
    #drop-zone {
      border:1.5px dashed var(--s300); border-radius:16px; background:#fff;
      padding:28px 24px; cursor:pointer; transition:border-color .15s, background .15s;
    }
    #drop-zone:hover  { border-color:var(--bd-500); background:var(--bd-50); }
    #drop-zone.active { border-color:var(--bd-500)!important; border-style:solid!important; background:var(--bd-50)!important; }
    .drop-arrow { transition:transform .2s cubic-bezier(0,0,.2,1); }
    #drop-zone.active .drop-arrow { transform:translateY(-3px); }

    /* ── Language chips ── */
    .lang-chip {
      display:flex; flex-direction:column; align-items:center; gap:5px;
      padding:10px 6px; border-radius:8px; border:1.5px solid var(--s300);
      background:#fff; cursor:pointer; transition:border-color .15s, background .15s;
      user-select:none; position:relative;
    }
    .lang-chip:hover { border-color:var(--s500); background:var(--s50); }
    .lang-chip.lch-on { border-color:var(--bd-500); background:var(--bd-50); }
    .lang-chip .chip-flag  { font-size:18px; line-height:1; }
    .lang-chip .chip-label { font-size:10px; font-weight:500; color:var(--s600); letter-spacing:.04em; text-transform:uppercase; }
    .lang-chip.lch-on .chip-label { color:var(--bd-700); }
    .lang-chip .chip-check {
      position:absolute; top:4px; right:4px; width:13px; height:13px; border-radius:50%;
      background:var(--bd-500); display:flex; align-items:center; justify-content:center;
      opacity:0; transition:opacity .15s;
    }
    .lang-chip.lch-on .chip-check { opacity:1; }

    /* ── Toggle ── */
    .tog-track {
      width:36px; height:20px; border-radius:10px; background:var(--s300);
      position:relative; cursor:pointer; transition:background .15s; flex-shrink:0;
    }
    .tog-track.tog-on { background:var(--bd-500); }
    .tog-thumb {
      position:absolute; top:3px; left:3px; width:14px; height:14px; border-radius:50%;
      background:#fff; transition:transform .15s cubic-bezier(.34,1.56,.64,1);
      box-shadow:0 1px 2px rgba(0,0,0,.2);
    }
    .tog-track.tog-on .tog-thumb { transform:translateX(16px); }

    /* ── Language progress cards ── */
    .lang-card {
      background:#fff; border:1px solid var(--border); border-radius:8px;
      padding:10px 4px; display:flex; flex-direction:column; align-items:center;
      gap:4px; transition:border-color .2s, background .2s;
    }
    .lang-card.lc-active { border-color:var(--bd-500); background:var(--bd-50); }
    .lang-card.lc-done   { border-color:var(--green-200); background:var(--green-50); }
    .lang-flag { font-size:18px; line-height:1; }
    .lang-abbr { font-size:10px; font-weight:500; color:var(--s600); letter-spacing:.04em; text-transform:uppercase; }
    .lang-card.lc-active .lang-abbr { color:var(--bd-700); }
    .lang-card.lc-done   .lang-abbr { color:var(--green-700); }
    .lang-dot { width:5px; height:5px; border-radius:50%; background:var(--s300); transition:background .2s; }
    .lang-card.lc-active .lang-dot { background:var(--bd-500); animation:pulse-dot 1.4s ease-in-out infinite; }
    .lang-card.lc-done   .lang-dot { background:var(--green-500); }

    /* ── Progress bar ── */
    #progress-bar { transition:width .7s cubic-bezier(.4,0,.2,1); }

    /* ── Insight card ── */
    #insight-wrap { transition:opacity .25s ease; }

    /* ── Buttons ── */
    .btn-ink {
      background:var(--bd-500); color:#fff; border-radius:8px; cursor:pointer;
      border:none; font-family:var(--font-text); font-weight:500; font-size:14px;
      transition:background .15s; letter-spacing:-.01em;
    }
    .btn-ink:hover:not(:disabled)  { background:var(--bd-600); }
    .btn-ink:active:not(:disabled) { background:var(--bd-700); }
    .btn-ink:disabled { opacity:.4; cursor:not-allowed; }

    .btn-ghost {
      border:1.5px solid var(--border); color:var(--s900); background:#fff;
      border-radius:8px; cursor:pointer; font-family:var(--font-text); font-weight:500; font-size:14px;
      transition:background .15s, border-color .15s; letter-spacing:-.01em;
    }
    .btn-ghost:hover { background:var(--s50); border-color:var(--s500); }
    .btn-ghost:disabled { opacity:.4; cursor:not-allowed; }

    /* ── Table ── */
    #review-thead th { background:var(--s50); border-bottom:1px solid var(--border); position:sticky; top:0; z-index:1; }
    #review-tbody tr:hover { background:var(--s50); }
  </style>
</head>
<body>

<!-- ── Header ── -->
<header style="background:#fff; border-bottom:1px solid var(--border); position:sticky; top:0; z-index:10;">
  <div style="max-width:720px; margin:0 auto; padding:0 24px; height:56px; display:flex; align-items:center; justify-content:space-between;">
    <div id="logo-home" style="display:flex; align-items:center; gap:10px; cursor:pointer; user-select:none;" title="Home">
      <div style="width:28px; height:28px; background:var(--s900); border-radius:6px; display:flex; align-items:center; justify-content:center; flex-shrink:0;">
        <span style="color:#fff; font-family:var(--font-display); font-size:14px; font-weight:700; letter-spacing:-.04em;">M</span>
      </div>
      <span style="font-family:var(--font-display); font-size:16px; font-weight:700; color:var(--s900); letter-spacing:-.02em;">Monotype</span>
    </div>
    <span style="font-size:11px; color:var(--s600);">Translation crew</span>
  </div>
</header>

<main style="max-width:720px; margin:0 auto; padding:40px 24px 80px;">

  <!-- ── Step indicator ── -->
  <div id="step-indicator" class="step-indicator" style="display:none;">
    <div class="si-step si-active" data-step="1"><div class="si-dot"></div><span class="si-label">Upload</span></div>
    <div class="si-line"></div>
    <div class="si-step" data-step="2"><div class="si-dot"></div><span class="si-label">Configure</span></div>
    <div class="si-line"></div>
    <div class="si-step" data-step="3"><div class="si-dot"></div><span class="si-label">Translating</span></div>
    <div class="si-line"></div>
    <div class="si-step" data-step="4"><div class="si-dot"></div><span class="si-label">Download</span></div>
  </div>

  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <!-- VIEW 1: UPLOAD                                                          -->
  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <div id="view-upload" class="anim-fade-up">

    <h1 style="font-family:var(--font-display); font-size:33px; line-height:40px; font-weight:700; letter-spacing:-.02em; color:var(--s900); margin-bottom:8px;">
      Translate your UI strings
    </h1>
    <p style="font-size:14px; color:var(--s700); line-height:1.6; max-width:520px; margin-bottom:28px;">
      Upload a file containing English source strings. We translate into five languages, guided by Monotype brand standards and approved terminology.
    </p>

    <!-- Drop zone -->
    <div id="drop-zone">
      <div style="display:flex; align-items:center; gap:20px;">
        <div class="drop-arrow" style="flex-shrink:0; width:48px; height:48px; border-radius:12px;
             background:var(--bd-50); border:1px solid var(--bd-200);
             display:flex; align-items:center; justify-content:center;">
          <svg width="20" height="20" viewBox="0 0 20 20" fill="none">
            <path d="M10 13V4M10 4L7 7M10 4L13 7" stroke="#1a73e8" stroke-width="1.6" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M3 14v1.5A1.5 1.5 0 004.5 17h11a1.5 1.5 0 001.5-1.5V14" stroke="#1a73e8" stroke-width="1.6" stroke-linecap="round"/>
          </svg>
        </div>
        <div>
          <p style="font-size:14px; font-weight:500; color:var(--s900); margin-bottom:4px;">
            Drop your file here — or <button id="browse-btn" style="color:var(--bd-500); text-decoration:underline; text-underline-offset:3px; background:none; border:none; cursor:pointer; font-size:14px; font-weight:500; padding:0; font-family:var(--font-text);">browse to upload</button>
          </p>
          <div style="display:flex; gap:4px; flex-wrap:wrap; margin-top:6px;">
            <span style="font-size:11px; padding:2px 8px; border-radius:4px; background:var(--s100); border:1px solid var(--border); color:var(--s700);">.xlsx</span>
            <span style="font-size:11px; padding:2px 8px; border-radius:4px; background:var(--s100); border:1px solid var(--border); color:var(--s700);">.pdf</span>
            <span style="font-size:11px; padding:2px 8px; border-radius:4px; background:var(--s100); border:1px solid var(--border); color:var(--s700);">.docx</span>
            <span style="font-size:11px; padding:2px 8px; border-radius:4px; background:var(--s100); color:var(--s600);">Max 10 MB</span>
          </div>
        </div>
      </div>
    </div>
    <input type="file" id="file-input" accept=".xlsx,.pdf,.docx,.doc" style="display:none;">

    <!-- File selected card -->
    <div id="file-card" style="display:none; margin-top:8px; padding:12px 16px; border-radius:8px;
         background:var(--s50); border:1px solid var(--border); align-items:center; gap:12px;">
      <div style="width:32px; height:32px; border-radius:6px; background:#fff; border:1px solid var(--border);
           display:flex; align-items:center; justify-content:center; flex-shrink:0;">
        <svg width="14" height="14" viewBox="0 0 14 14" fill="none">
          <path d="M8 1H3a1 1 0 00-1 1v10a1 1 0 001 1h8a1 1 0 001-1V6L8 1z" stroke="var(--s600)" stroke-width="1.2"/>
          <path d="M8 1v5h5" stroke="var(--s600)" stroke-width="1.2" stroke-linecap="round"/>
        </svg>
      </div>
      <div style="flex:1; min-width:0;">
        <p id="file-name" style="font-size:13px; font-weight:500; color:var(--s900); overflow:hidden; text-overflow:ellipsis; white-space:nowrap;"></p>
        <div style="display:flex; align-items:center; gap:8px; margin-top:3px; flex-wrap:wrap;">
          <p id="file-size" style="font-size:11px; color:var(--s600);"></p>
          <div id="preview-spinner" style="display:none; width:10px; height:10px; border-radius:50%;
               border:1.5px solid var(--border); border-top-color:var(--s600); animation:spin .7s linear infinite;"></div>
          <div id="string-count-pill" style="display:none; font-size:11px; font-weight:500; padding:2px 8px;
               border-radius:9999px; background:var(--bd-100); color:var(--bd-700); border:1px solid var(--bd-200);"></div>
        </div>
      </div>
      <button id="remove-file" style="width:20px; height:20px; border-radius:50%; border:none; background:none;
              cursor:pointer; color:var(--s500); display:flex; align-items:center; justify-content:center; padding:0;">
        <svg width="10" height="10" viewBox="0 0 10 10" fill="none">
          <path d="M1 1l8 8M9 1L1 9" stroke="currentColor" stroke-width="1.5" stroke-linecap="round"/>
        </svg>
      </button>
    </div>

    <!-- Trust signals -->
    <div style="display:flex; align-items:center; gap:6px; margin-top:16px; flex-wrap:wrap;">
      <div style="display:flex; align-items:center; gap:5px; padding:4px 10px; border-radius:9999px; background:var(--s100); border:1px solid var(--border);">
        <div style="width:5px; height:5px; border-radius:50%; background:var(--s600); flex-shrink:0;"></div>
        <span style="font-size:11px; color:var(--s700);">Brand-aligned terminology</span>
      </div>
      <div style="display:flex; align-items:center; gap:5px; padding:4px 10px; border-radius:9999px; background:var(--s100); border:1px solid var(--border);">
        <div style="width:5px; height:5px; border-radius:50%; background:var(--s600); flex-shrink:0;"></div>
        <span style="font-size:11px; color:var(--s700);">Preserves tone &amp; UI length</span>
      </div>
      <div style="display:flex; align-items:center; gap:5px; padding:4px 10px; border-radius:9999px; background:var(--s100); border:1px solid var(--border);">
        <div style="width:5px; height:5px; border-radius:50%; background:var(--s600); flex-shrink:0;"></div>
        <span style="font-size:11px; color:var(--s700);">5 languages at once</span>
      </div>
    </div>

    <!-- CTA -->
    <button id="continue-btn" disabled class="btn-ink"
            style="width:100%; padding:12px; font-size:14px; margin-top:20px;
                   display:flex; align-items:center; justify-content:center; gap:8px;">
      <span>Continue</span>
      <svg width="14" height="14" viewBox="0 0 14 14" fill="none">
        <path d="M3 7h8M8 4l3 3-3 3" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
      </svg>
    </button>
  </div>

  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <!-- VIEW 2: CONFIGURE                                                       -->
  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <div id="view-configure" style="display:none;" class="anim-fade-up">

    <!-- File summary -->
    <div style="display:flex; align-items:center; gap:12px; padding:12px 16px; background:#fff;
         border:1px solid var(--border); border-radius:8px; margin-bottom:24px;">
      <div style="width:32px; height:32px; border-radius:6px; background:var(--s100); flex-shrink:0;
           display:flex; align-items:center; justify-content:center;">
        <svg width="14" height="14" viewBox="0 0 14 14" fill="none">
          <path d="M8 1H3a1 1 0 00-1 1v10a1 1 0 001 1h8a1 1 0 001-1V6L8 1z" stroke="var(--s600)" stroke-width="1.2"/>
          <path d="M8 1v5h5" stroke="var(--s600)" stroke-width="1.2" stroke-linecap="round"/>
        </svg>
      </div>
      <div style="flex:1; min-width:0;">
        <p id="conf-filename" style="font-size:13px; font-weight:500; color:var(--s900); overflow:hidden; text-overflow:ellipsis; white-space:nowrap;"></p>
        <p id="conf-meta" style="font-size:11px; color:var(--s600); margin-top:2px;"></p>
      </div>
      <button id="conf-change" style="font-size:12px; color:var(--bd-500); background:none; border:none; cursor:pointer;
              font-weight:500; white-space:nowrap; flex-shrink:0; font-family:var(--font-text);">
        Change file
      </button>
    </div>

    <!-- Language selection -->
    <div style="background:#fff; border:1px solid var(--border); border-radius:16px; padding:20px; margin-bottom:12px;">
      <p style="font-size:11px; font-weight:500; color:var(--s600); letter-spacing:.04em; text-transform:uppercase; margin-bottom:4px;">Target languages</p>
      <p style="font-size:13px; color:var(--s700); margin-bottom:14px;">All five are selected. Tap to deselect any you don't need.</p>
      <div id="lang-chips" style="display:grid; grid-template-columns:repeat(5,1fr); gap:8px;"></div>
    </div>

    <!-- Length toggle -->
    <div style="padding:16px 20px; background:#fff; border:1px solid var(--border); border-radius:16px; margin-bottom:24px;">
      <div style="display:flex; align-items:center; justify-content:space-between; gap:16px;">
        <div>
          <p style="font-size:14px; font-weight:500; color:var(--s900); margin-bottom:3px;">Optimise for UI length</p>
          <p style="font-size:12px; color:var(--s600); line-height:1.5;">Keep translations concise — helps strings fit buttons and labels at smaller sizes.</p>
        </div>
        <div id="length-toggle" class="tog-track tog-on" onclick="toggleLength()">
          <div class="tog-thumb"></div>
        </div>
      </div>
    </div>

    <!-- Start button -->
    <button id="translate-btn" class="btn-ink"
            style="width:100%; padding:12px; font-size:14px;
                   display:flex; align-items:center; justify-content:center; gap:8px;">
      <span id="translate-btn-label">Start translation</span>
      <span id="translate-btn-icon" style="display:flex; align-items:center;">
        <svg width="14" height="14" viewBox="0 0 14 14" fill="none">
          <path d="M3 7h8M8 4l3 3-3 3" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>
      </span>
    </button>

    <p style="font-size:11px; color:var(--s600); text-align:center; margin-top:12px; line-height:1.6;">
      Translation takes 3–8 minutes. Progress shown below — you can leave this tab open.
    </p>
  </div>

  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <!-- VIEW 3: PROCESSING                                                      -->
  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <div id="view-processing" style="display:none;" class="anim-fade-up">

    <!-- File + elapsed -->
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:16px;">
      <div style="display:flex; align-items:center; gap:8px;">
        <svg width="12" height="14" viewBox="0 0 12 14" fill="none">
          <path d="M7 1H2a1 1 0 00-1 1v10a1 1 0 001 1h8a1 1 0 001-1V5L7 1z" stroke="var(--s500)" stroke-width="1.2"/>
          <path d="M7 1v4h4" stroke="var(--s500)" stroke-width="1.2" stroke-linecap="round"/>
        </svg>
        <span id="proc-filename" style="font-size:12px; font-weight:500; color:var(--s700);"></span>
      </div>
      <span id="elapsed-disp" style="font-size:12px; color:var(--s600); font-variant-numeric:tabular-nums;"></span>
    </div>

    <!-- Progress bar -->
    <div style="margin-bottom:16px;">
      <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:6px;">
        <span id="progress-stage-label" style="font-size:12px; color:var(--s700); font-weight:500;">Initialising…</span>
        <span id="progress-pct" style="font-size:12px; color:var(--s600); font-variant-numeric:tabular-nums;">0%</span>
      </div>
      <div style="height:4px; background:var(--s200); border-radius:9999px; overflow:hidden;">
        <div id="progress-bar" style="height:100%; background:var(--bd-500); width:2%; border-radius:9999px;"></div>
      </div>
    </div>

    <!-- Active task banner -->
    <div id="active-task-card" style="background:var(--s900); border-radius:12px; padding:14px 18px;
         display:flex; align-items:center; gap:14px; margin-bottom:16px;">
      <div style="width:7px; height:7px; border-radius:50%; background:#fff; flex-shrink:0;
           animation:pulse-dot 1.4s ease-in-out infinite;"></div>
      <div style="min-width:0;">
        <p id="active-task-label" style="font-size:13px; font-weight:500; color:#fff; margin-bottom:2px;">Initialising…</p>
        <p id="active-task-desc"  style="font-size:11px; color:rgba(255,255,255,.45); white-space:nowrap; overflow:hidden; text-overflow:ellipsis;"></p>
      </div>
    </div>

    <!-- Language progress cards -->
    <div style="display:grid; grid-template-columns:repeat(5,1fr); gap:8px; margin-bottom:24px;">
      <div id="lang-fr" class="lang-card"><span class="lang-flag">🇫🇷</span><span class="lang-abbr">FR</span><div class="lang-dot"></div></div>
      <div id="lang-de" class="lang-card"><span class="lang-flag">🇩🇪</span><span class="lang-abbr">DE</span><div class="lang-dot"></div></div>
      <div id="lang-pt" class="lang-card"><span class="lang-flag">🇧🇷</span><span class="lang-abbr">PT</span><div class="lang-dot"></div></div>
      <div id="lang-ja" class="lang-card"><span class="lang-flag">🇯🇵</span><span class="lang-abbr">JA</span><div class="lang-dot"></div></div>
      <div id="lang-es" class="lang-card"><span class="lang-flag">🇪🇸</span><span class="lang-abbr">ES</span><div class="lang-dot"></div></div>
    </div>

    <!-- Insight card -->
    <div>
      <p style="font-size:10px; color:var(--s600); letter-spacing:.08em; text-transform:uppercase; margin-bottom:10px; font-weight:500;">Did you know</p>
      <div id="insight-wrap" style="background:#fff; border:1px solid var(--border); border-radius:12px; padding:18px;">
        <p id="insight-title" style="font-size:13px; font-weight:500; color:var(--s900); margin-bottom:6px;"></p>
        <p id="insight-body"  style="font-size:12px; line-height:1.75; color:var(--s700);"></p>
        <div id="insight-dots" style="display:flex; gap:4px; margin-top:14px;"></div>
      </div>
    </div>

    <p style="font-size:11px; color:var(--s600); text-align:center; margin-top:24px; line-height:1.6;">
      Your translations are getting ready — you can leave this tab open.
    </p>
    <p id="job-id-disp" style="font-size:10px; color:var(--s300); text-align:center; font-family:var(--font-mono); margin-top:4px;"></p>
    <div style="text-align:center; margin-top:16px;">
      <button id="cancel-btn" class="btn-ghost" style="padding:8px 28px; font-size:12px;">
        Cancel translation
      </button>
    </div>
  </div>

  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <!-- VIEW 4: RESULTS                                                         -->
  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <div id="view-results" style="display:none;">

    <!-- Headline -->
    <div style="margin-bottom:24px;" class="anim-slide-up">
      <div style="display:flex; align-items:center; gap:8px; margin-bottom:8px;">
        <div style="width:24px; height:24px; border-radius:50%; background:var(--green-50); border:1px solid var(--green-200);
             display:flex; align-items:center; justify-content:center; flex-shrink:0;">
          <svg width="12" height="12" viewBox="0 0 12 12" fill="none">
            <path d="M2 6.5l2.5 2.5 5.5-5.5" stroke="#26a568" stroke-width="1.6"
                  stroke-linecap="round" stroke-linejoin="round"
                  stroke-dasharray="12" stroke-dashoffset="12"
                  style="animation:draw-check .35s .05s ease forwards;"/>
          </svg>
        </div>
        <span style="font-size:11px; font-weight:500; color:var(--green-700); letter-spacing:.06em; text-transform:uppercase;">Complete</span>
      </div>
      <h2 style="font-family:var(--font-display); font-size:33px; line-height:40px; font-weight:700; letter-spacing:-.02em; color:var(--s900); margin-bottom:6px;">Ready to download.</h2>
      <p id="result-summary" style="font-size:13px; color:var(--s700);"></p>
    </div>

    <!-- Language completion pills -->
    <div id="lang-pills" style="display:flex; gap:6px; flex-wrap:wrap; margin-bottom:16px;"></div>

    <!-- Download card -->
    <div style="background:var(--s900); border-radius:16px; padding:18px 20px; display:flex;
         align-items:center; justify-content:space-between; gap:16px; margin-bottom:16px;">
      <div>
        <p id="dl-format-label" style="font-size:13px; font-weight:500; color:#fff; margin-bottom:4px;">All translations · Excel</p>
        <p id="dl-filename" style="font-size:11px; color:rgba(255,255,255,.4);"></p>
      </div>
      <button id="download-btn"
              style="display:flex; align-items:center; gap:8px; padding:8px 16px; border-radius:8px;
                     background:rgba(255,255,255,.1); border:1px solid rgba(255,255,255,.15);
                     color:#fff; font-size:13px; font-weight:500; cursor:pointer; white-space:nowrap;
                     font-family:var(--font-text); letter-spacing:-.01em; transition:background .15s;"
              onmouseover="this.style.background='rgba(255,255,255,.18)'"
              onmouseout="this.style.background='rgba(255,255,255,.1)'">
        <svg width="13" height="13" viewBox="0 0 13 13" fill="none">
          <path d="M6.5 1v8M3.5 6l3 3 3-3" stroke="white" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
          <path d="M1 10v1a1.5 1.5 0 001.5 1.5h8A1.5 1.5 0 0012 11v-1" stroke="white" stroke-width="1.5" stroke-linecap="round"/>
        </svg>
        Download
      </button>
    </div>

    <!-- Preview table -->
    <div style="background:#fff; border:1px solid var(--border); border-radius:16px; overflow:hidden; margin-bottom:16px;">
      <div style="padding:14px 20px; border-bottom:1px solid var(--border); background:var(--s50);
           display:flex; align-items:center; justify-content:space-between; gap:12px;">
        <p style="font-size:11px; font-weight:500; color:var(--s600); letter-spacing:.06em; text-transform:uppercase; white-space:nowrap;">
          Preview translations
        </p>
        <input id="search-input" type="text" placeholder="Search strings…"
               style="font-size:13px; padding:6px 12px; border-radius:8px; border:1px solid var(--border);
                      background:#fff; color:var(--s900); width:180px; outline:none; font-family:var(--font-text);
                      transition:border-color .15s;"
               onfocus="this.style.borderColor='var(--bd-500)'" onblur="this.style.borderColor='var(--border)'"/>
      </div>
      <div style="overflow-x:auto; max-height:420px; overflow-y:auto;">
        <table style="width:100%; border-collapse:collapse; font-size:13px;">
          <thead id="review-thead"></thead>
          <tbody id="review-tbody"></tbody>
        </table>
      </div>
      <p id="no-matches" style="display:none; text-align:center; font-size:13px; color:var(--s600); padding:24px;">No matching strings.</p>
    </div>

    <button id="btn-new" class="btn-ghost" style="width:100%; padding:11px; font-size:13px;">
      Translate another file
    </button>
  </div>

  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <!-- VIEW: ERROR                                                              -->
  <!-- ═══════════════════════════════════════════════════════════════════════ -->
  <div id="view-error" style="display:none; text-align:center; padding:48px 0;" class="anim-fade-up">
    <div style="width:44px; height:44px; border-radius:50%; background:var(--red-50); border:1px solid var(--red-200);
         display:inline-flex; align-items:center; justify-content:center; margin-bottom:16px;">
      <svg width="18" height="18" viewBox="0 0 18 18" fill="none">
        <path d="M9 6v4M9 12.5h.01M9 1L1 16h16L9 1z" stroke="#dc0024" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
      </svg>
    </div>
    <h3 style="font-family:var(--font-display); font-size:23px; line-height:32px; font-weight:700; letter-spacing:-.01em; color:var(--s900); margin-bottom:8px;">
      Something went wrong.
    </h3>
    <p id="error-text" style="font-size:13px; color:var(--s700); max-width:380px; margin:0 auto 24px; line-height:1.6;"></p>
    <button id="btn-retry" class="btn-ink" style="padding:10px 28px; font-size:13px;">
      Try again
    </button>
  </div>

</main>

<script>
// ── Data ──────────────────────────────────────────────────────────────────────
const LANG_LABELS = {
  en:'English', fr:'French', de:'German',
  pt_BR:'Portuguese (pt-BR)', ja:'Japanese', es_ES:'Spanish (es-ES)'
};

const LANG_CONFIG = [
  { id:'fr', flag:'🇫🇷', name:'French',     abbr:'FR' },
  { id:'de', flag:'🇩🇪', name:'German',     abbr:'DE' },
  { id:'pt', flag:'🇧🇷', name:'Portuguese', abbr:'PT' },
  { id:'ja', flag:'🇯🇵', name:'Japanese',   abbr:'JA' },
  { id:'es', flag:'🇪🇸', name:'Spanish (es-ES)', abbr:'ES' },
];

const PIPELINE_STAGES = [
  { id:'parse',   label:'Reading source strings',                desc:'Parsing your file and counting UI strings to translate',        est:0   },
  { id:'brand',   label:'Loading Monotype brand voice',          desc:'Ingesting glossary, brand guidelines and approved terminology',  est:8   },
  { id:'context', label:'Mapping linguistic context',            desc:'Building translation memory and identifying edge cases',        est:22  },
  { id:'fr',      label:'Crafting French copy',                  desc:'Translating with idiomatic French phrasing and brand tone',     est:70  },
  { id:'de',      label:'Structuring German grammar',            desc:'Handling compound nouns, formal register and UI string length', est:120 },
  { id:'pt',      label:'Adapting for Brazilian Portuguese',     desc:'Applying pt-BR orthography, vocabulary and article gender',     est:170 },
  { id:'ja',      label:'Rendering Japanese scripts',            desc:'Balancing Kanji, Katakana and natural particle selection',      est:220 },
  { id:'es',      label:'Localising for Spain Spanish', desc:'Targeting es-ES (Castilian) register, vocabulary and cultural context',    est:270 },
  { id:'review',  label:'Reviewing all 5 languages',             desc:'Cross-checking glossary adherence, tone and brand consistency', est:320 },
  { id:'report',  label:'Writing production report',             desc:'Summarising translation decisions and flagging edge cases',     est:430 },
];

const INSIGHTS = [
  { title:'Why Japanese UI copy is shorter',
    body:'Japanese strings run 40–60% shorter than English — not because less is said, but because Kanji packs meaning densely. A single character like 設定 means "Settings".' },
  { title:'The German button problem',
    body:'German UI strings average 30% longer than English. "Einstellungen" fits fine, but "Benachrichtigungseinstellungen" (notification settings) can overflow a 320 px button.' },
  { title:'French punctuation is strict',
    body:'A non-breaking space before : ; ? and ! is mandatory in French typography — not optional. Missing it immediately signals a non-native translation to French readers.' },
  { title:'Two very different Portugals',
    body:'Brazilian (pt-BR) and European Portuguese diverge in vocabulary, verb forms, and orthography. Many product terms differ entirely between the two variants.' },
  { title:'は vs が — the particle problem',
    body:'In Japanese, は marks the topic and が marks the subject with emphasis. Getting this wrong does not break meaning, but reads as unnatural to any native speaker.' },
  { title:'es-ES is Castilian Spanish',
    body:'Using Spain Spanish (Castilian). Vocabulary and register are aligned with Spain — distinct from Latin American Spanish (es-419).' },
  { title:'Brand voice travels across languages',
    body:'Good localisation recreates emotional register, not just vocabulary. A premium product must sound premium in Tokyo and Berlin — which requires cultural adaptation, not literal translation.' },
  { title:"Monotype's 150,000-font library",
    body:"Monotype's type library spans centuries — from Garamond in the 1530s to modern variable fonts. The UI copy across all products needs to feel as considered as the typefaces themselves." },
];

// ── State ─────────────────────────────────────────────────────────────────────
let selectedFile   = null;
let selectedFileType = 'xlsx';
let selectedLangs  = new Set(['fr','de','pt','ja','es']);
let selectedTone   = 'informal';
let optimizeLen    = true;
let stringCount    = null;
let currentJobId   = null;
let currentStep    = 0;
let pollTimer      = null;
let elapsedTimer   = null;
let insightTimer   = null;
let startedAt      = null;
let insightIdx     = 0;
let allRows        = [];

// ── Helpers ───────────────────────────────────────────────────────────────────
const $ = id => document.getElementById(id);
function showEl(id, flex) { const e=$(id); if(e) e.style.display = flex?'flex':'block'; }
function hideEl(id)        { const e=$(id); if(e) e.style.display = 'none'; }
function esc(v) { return String(v??'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;'); }
function fmtBytes(b) { return b<1024?b+' B':b<1048576?(b/1024).toFixed(1)+' KB':(b/1048576).toFixed(1)+' MB'; }

// ── Step navigation ───────────────────────────────────────────────────────────
const STEP_VIEWS = [null,'upload','configure','processing','results'];

function switchView(name) {
  ['upload','configure','processing','results','error'].forEach(v => hideEl('view-'+v));
  const el = $('view-'+name);
  el.style.display = 'block';
  el.classList.remove('anim-fade-up','anim-slide-up');
  void el.offsetWidth;
  el.classList.add('anim-fade-up');
}

function goToStep(n) {
  currentStep = n;
  const ind = $('step-indicator');
  ind.style.display = (n >= 1 && n <= 4) ? 'flex' : 'none';
  document.querySelectorAll('.si-step').forEach(s => {
    const sn = parseInt(s.dataset.step);
    s.classList.remove('si-active','si-done');
    if (sn === n) s.classList.add('si-active');
    else if (sn < n) s.classList.add('si-done');
  });
  document.querySelectorAll('.si-line').forEach((ln, i) => {
    ln.style.background = (i + 1 < n) ? 'var(--muted)' : 'var(--border)';
  });
  if (STEP_VIEWS[n]) switchView(STEP_VIEWS[n]);
}

// ── File handling ─────────────────────────────────────────────────────────────
const dropZone = $('drop-zone');
dropZone.addEventListener('dragover',  e => { e.preventDefault(); dropZone.classList.add('active'); });
dropZone.addEventListener('dragleave', () => dropZone.classList.remove('active'));
dropZone.addEventListener('drop', e => {
  e.preventDefault(); dropZone.classList.remove('active');
  if (e.dataTransfer.files[0]) handleFile(e.dataTransfer.files[0]);
});
dropZone.addEventListener('click', e => { if (e.target.id !== 'browse-btn') $('file-input').click(); });
$('browse-btn').addEventListener('click', e => { e.stopPropagation(); $('file-input').click(); });
$('file-input').addEventListener('change', () => { if ($('file-input').files[0]) handleFile($('file-input').files[0]); });
$('remove-file').addEventListener('click', clearFile);

async function handleFile(file) {
  if (!/\.(xlsx|pdf|docx|doc)$/i.test(file.name)) { alert('Please upload an .xlsx, .pdf, or .docx file.'); return; }
  if (file.size > 10*1024*1024) { alert('File must be smaller than 10 MB.'); return; }
  selectedFile = file;
  selectedFileType = /\.(docx|doc)$/i.test(file.name) ? 'docx' : 'xlsx';
  $('file-name').textContent = file.name;
  $('file-size').textContent = fmtBytes(file.size);
  hideEl('string-count-pill'); hideEl('drop-zone');
  showEl('preview-spinner'); showEl('file-card', true);
  $('continue-btn').disabled = true;

  try {
    const fd = new FormData(); fd.append('file', file);
    const res = await fetch('/api/preview', { method:'POST', body:fd });
    if (res.ok) {
      const data = await res.json();
      stringCount = data.string_count;
      hideEl('preview-spinner');
      if (stringCount) {
        const pill = $('string-count-pill');
        pill.textContent = stringCount + ' strings detected';
        showEl('string-count-pill');
        pill.style.animation = 'none'; void pill.offsetWidth;
        pill.style.animation = 'pop .35s cubic-bezier(.16,1,.3,1) both';
      }
    } else { hideEl('preview-spinner'); }
  } catch(_) { hideEl('preview-spinner'); }
  $('continue-btn').disabled = false;
}

function clearFile() {
  selectedFile = null; stringCount = null; $('file-input').value = '';
  hideEl('file-card'); showEl('drop-zone'); $('continue-btn').disabled = true;
}

// ── Step 1 → Step 2 ───────────────────────────────────────────────────────────
$('continue-btn').addEventListener('click', () => {
  if (!selectedFile) return;
  $('conf-filename').textContent = selectedFile.name;
  $('conf-meta').textContent = fmtBytes(selectedFile.size) + (stringCount ? ' · ' + stringCount + ' strings detected' : '');
  buildLangChips();
  goToStep(2);
});
$('conf-change').addEventListener('click', () => goToStep(1));

// ── Language chips ────────────────────────────────────────────────────────────
function buildLangChips() {
  $('lang-chips').innerHTML = LANG_CONFIG.map(l => `
    <div class="lang-chip${selectedLangs.has(l.id) ? ' lch-on' : ''}" data-lang="${l.id}">
      <span class="chip-flag">${l.flag}</span>
      <span class="chip-label">${l.abbr}</span>
      <div class="chip-check">
        <svg width="8" height="8" viewBox="0 0 8 8" fill="none">
          <path d="M1 4.5l2 2 4-4" stroke="white" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>
      </div>
    </div>`).join('');
  $('lang-chips').querySelectorAll('.lang-chip').forEach(chip => {
    chip.addEventListener('click', () => {
      const lang = chip.dataset.lang;
      if (selectedLangs.has(lang) && selectedLangs.size <= 1) return;
      if (selectedLangs.has(lang)) { selectedLangs.delete(lang); chip.classList.remove('lch-on'); }
      else { selectedLangs.add(lang); chip.classList.add('lch-on'); }
    });
  });
}

// ── Length toggle ─────────────────────────────────────────────────────────────
function toggleLength() {
  optimizeLen = !optimizeLen;
  optimizeLen ? $('length-toggle').classList.add('tog-on') : $('length-toggle').classList.remove('tog-on');
}

// ── Start translation ─────────────────────────────────────────────────────────
$('translate-btn').addEventListener('click', startTranslation);

async function startTranslation() {
  if (!selectedFile) return;
  const btn = $('translate-btn');
  btn.disabled = true;
  $('translate-btn-label').textContent = 'Starting…';
  $('translate-btn-icon').innerHTML = `
    <svg width="14" height="14" viewBox="0 0 14 14" fill="none" style="animation:spin .7s linear infinite;">
      <circle cx="7" cy="7" r="5" stroke="rgba(255,255,255,.3)" stroke-width="1.5"/>
      <path d="M7 2a5 5 0 015 5" stroke="white" stroke-width="1.5" stroke-linecap="round"/>
    </svg>`;

  const fd = new FormData();
  fd.append('file', selectedFile);
  fd.append('tone', selectedTone);
  fd.append('languages', Array.from(selectedLangs).join(','));
  fd.append('optimize_length', String(optimizeLen));

  try {
    const res = await fetch('/api/translate', { method:'POST', body:fd });
    if (!res.ok) throw new Error(await res.text());
    const data = await res.json();
    currentJobId = data.job_id;
    startedAt    = Date.now();
    $('proc-filename').textContent = selectedFile.name;
    $('job-id-disp').textContent   = 'job ' + currentJobId;
    goToStep(3);
    initProcessingView();
    pollTimer = setInterval(pollStatus, 5000);
    pollStatus();
  } catch(err) {
    btn.disabled = false;
    $('translate-btn-label').textContent = 'Start Translation';
    $('translate-btn-icon').innerHTML = `<svg width="14" height="14" viewBox="0 0 14 14" fill="none">
      <path d="M3 7h8M8 4l3 3-3 3" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
    </svg>`;
    alert('Could not start: ' + err.message);
  }
}

// ── Processing view ───────────────────────────────────────────────────────────
function initProcessingView() {
  $('active-task-label').textContent = PIPELINE_STAGES[0].label;
  $('active-task-desc').textContent  = PIPELINE_STAGES[0].desc;
  ['fr','de','pt','ja','es'].forEach(l => { const e=$('lang-'+l); if(e) e.className='lang-card'; });

  $('insight-dots').innerHTML = INSIGHTS.map((_,i) =>
    `<div data-dot="${i}" style="width:5px;height:5px;border-radius:50%;
          background:${i===0?'var(--ink)':'var(--border)'};transition:background .3s;"></div>`
  ).join('');
  showInsight(0);

  elapsedTimer = setInterval(() => {
    const sec = Math.floor((Date.now()-startedAt)/1000);
    const m = Math.floor(sec/60), s = sec%60;
    $('elapsed-disp').textContent = m + 'm ' + String(s).padStart(2,'0') + 's';
    advanceStages(sec);
  }, 1000);
  insightTimer = setInterval(() => showInsight(insightIdx+1), 9000);
}

const LANG_STAGE_IDX = { fr:3, de:4, pt:5, ja:6, es:7 };

function advanceStages(sec) {
  let ai = 0;
  PIPELINE_STAGES.forEach((s,i) => { if (sec >= s.est) ai = i; });
  Object.entries(LANG_STAGE_IDX).forEach(([l, si]) => {
    const e = $('lang-'+l); if (!e) return;
    e.className = ai > si ? 'lang-card lc-done' : ai === si ? 'lang-card lc-active' : 'lang-card';
  });
  const stage = PIPELINE_STAGES[ai];
  $('active-task-label').textContent = stage.label;
  $('active-task-desc').textContent  = stage.desc;
  const maxEst = PIPELINE_STAGES[PIPELINE_STAGES.length-1].est + 60;
  const pct = Math.min(95, 2 + (sec/maxEst)*93);
  $('progress-bar').style.width = pct + '%';
  $('progress-pct').textContent = Math.round(pct) + '%';
  $('progress-stage-label').textContent = stage.label;
}

function showInsight(idx) {
  insightIdx = ((idx % INSIGHTS.length) + INSIGHTS.length) % INSIGHTS.length;
  const wrap = $('insight-wrap');
  wrap.style.opacity = '0';
  setTimeout(() => {
    $('insight-title').textContent = INSIGHTS[insightIdx].title;
    $('insight-body').textContent  = INSIGHTS[insightIdx].body;
    document.querySelectorAll('[data-dot]').forEach(d => {
      d.style.background = parseInt(d.dataset.dot) === insightIdx ? 'var(--ink)' : 'var(--border)';
    });
    wrap.style.opacity = '1';
  }, 350);
}

// ── Poll status ───────────────────────────────────────────────────────────────
async function pollStatus() {
  try {
    const res = await fetch('/api/status/' + currentJobId);
    if (!res.ok) return;
    const data = await res.json();
    if (data.status === 'complete') {
      clearInterval(pollTimer); clearInterval(elapsedTimer); clearInterval(insightTimer);
      $('progress-bar').style.width  = '100%';
      $('progress-pct').textContent  = '100%';
      $('active-task-label').textContent = 'Translation complete';
      $('active-task-desc').textContent  = 'All languages reviewed and ready to download';
      ['fr','de','pt','ja','es'].forEach(l => { const e=$('lang-'+l); if(e) e.className='lang-card lc-done'; });
      setTimeout(() => { goToStep(4); showResults(data); }, 800);
    } else if (data.status === 'failed') {
      clearInterval(pollTimer); clearInterval(elapsedTimer); clearInterval(insightTimer);
      switchView('error');
      $('error-text').textContent = data.error || 'An unknown error occurred.';
    } else if (data.status === 'cancelled') {
      clearInterval(pollTimer); clearInterval(elapsedTimer); clearInterval(insightTimer);
      resetUI();
    }
  } catch(_) {}
}

// ── Results ───────────────────────────────────────────────────────────────────
function showResults(data) {
  const n = data.review_data ? data.review_data.length : '?';
  const isDocx = (data.file_type === 'docx') || selectedFileType === 'docx';
  const unit = isDocx ? 'segment' : 'string';
  $('result-summary').textContent = n + ' ' + unit + (n!==1?'s':'') + ' · ' + selectedLangs.size + ' language' + (selectedLangs.size !== 1 ? 's' : '');
  const baseName = selectedFile ? selectedFile.name.replace(/\.\w+$/, '') : 'file';
  $('dl-filename').textContent = baseName + (isDocx ? '_translated.zip' : '_translated.xlsx');
  $('dl-format-label').textContent = 'All translations · ' + (isDocx ? 'Zip' : 'Excel');

  $('lang-pills').innerHTML = LANG_CONFIG.filter(l => selectedLangs.has(l.id)).map(l =>
    `<div style="display:flex;align-items:center;gap:5px;padding:5px 11px;border-radius:20px;
          background:var(--green-bg);border:1px solid var(--green-bd);">
       <span style="font-size:13px;">${l.flag}</span>
       <span style="font-size:11px;font-weight:600;color:var(--green);">${l.name}</span>
       <svg width="10" height="10" viewBox="0 0 10 10" fill="none">
         <path d="M1.5 5.5l2.5 2.5 4.5-5" stroke="var(--green)" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"/>
       </svg>
     </div>`
  ).join('');

  if (data.review_data && data.review_data.length > 0) {
    allRows = data.review_data;
    renderTable(allRows);
  }
}

$('download-btn').addEventListener('click', () => { window.location.href = '/api/download/' + currentJobId; });

function renderTable(rows) {
  if (!rows.length) return;
  const keys = Object.keys(rows[0]).filter(k => k !== 'row_index');
  $('review-thead').innerHTML = '<tr>' + keys.map(k =>
    '<th style="padding:10px 16px;text-align:left;font-size:11px;font-weight:600;color:var(--muted);letter-spacing:.07em;text-transform:uppercase;white-space:nowrap;">' +
    esc(LANG_LABELS[k]||k) + '</th>'
  ).join('') + '</tr>';
  $('review-tbody').innerHTML = rows.map((row,i) =>
    '<tr style="background:'+(i%2?'var(--paper)':'white')+'">' +
    keys.map(k =>
      '<td style="padding:10px 16px;font-size:12px;color:var(--ink);max-width:200px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;vertical-align:top;" title="'+esc(row[k])+'">'+esc(row[k])+'</td>'
    ).join('') + '</tr>'
  ).join('');
  hideEl('no-matches');
}

let _searchTimer = null;
$('search-input').addEventListener('input', e => {
  clearTimeout(_searchTimer);
  _searchTimer = setTimeout(() => {
    const q = e.target.value.toLowerCase();
    const filtered = q ? allRows.filter(r => Object.values(r).some(v => String(v??'').toLowerCase().includes(q))) : allRows;
    renderTable(filtered);
    filtered.length ? hideEl('no-matches') : showEl('no-matches');
  }, 150);
});

// ── Cancel ────────────────────────────────────────────────────────────────────
$('cancel-btn').addEventListener('click', async () => {
  if (!currentJobId) return;
  const btn = $('cancel-btn');
  btn.disabled = true; btn.textContent = 'Cancelling…';
  try { await fetch('/api/cancel/' + currentJobId, { method:'DELETE' }); } catch(_) {}
  [pollTimer, elapsedTimer, insightTimer].forEach(t => t && clearInterval(t));
  resetUI();
});

// ── Logo → home ───────────────────────────────────────────────────────────────
$('logo-home').addEventListener('click', () => {
  if (currentJobId && !confirm('A translation is in progress. Cancel it and go back to the upload screen?')) return;
  if (currentJobId) fetch('/api/cancel/' + currentJobId, { method:'DELETE' }).catch(()=>{});
  [pollTimer, elapsedTimer, insightTimer].forEach(t => t && clearInterval(t));
  resetUI();
});

// ── Reset ─────────────────────────────────────────────────────────────────────
$('btn-new').addEventListener('click', resetUI);
$('btn-retry').addEventListener('click', resetUI);

function resetUI() {
  [pollTimer, elapsedTimer, insightTimer].forEach(t => t && clearInterval(t));
  currentJobId = null; selectedFile = null; startedAt = null;
  allRows = []; insightIdx = 0; stringCount = null;
  selectedLangs = new Set(['fr','de','pt','ja','es']);
  selectedTone  = 'informal';
  optimizeLen   = true;
  $('file-input').value = '';
  const si = $('search-input'); if (si) si.value = '';
  const cb = $('cancel-btn'); if (cb) { cb.disabled = false; cb.textContent = 'Cancel Translation'; }
  $('translate-btn').disabled = false;
  $('translate-btn-label').textContent = 'Start Translation';
  $('translate-btn-icon').innerHTML = `<svg width="14" height="14" viewBox="0 0 14 14" fill="none">
    <path d="M3 7h8M8 4l3 3-3 3" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/>
  </svg>`;
  $('length-toggle').classList.add('tog-on');
  hideEl('file-card'); showEl('drop-zone');
  $('continue-btn').disabled = true;
  goToStep(1);
}

// ── Boot ──────────────────────────────────────────────────────────────────────
goToStep(1);
</script>
</body>
</html>"""


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def index():
    return _HTML


_ALLOWED_EXTS = {".xlsx", ".pdf", ".docx", ".doc"}


@app.post("/api/preview")
async def preview_file(file: UploadFile = File(...)):
    """Return basic metadata (string count) without persisting the file."""
    fname = file.filename or ""
    ext   = Path(fname).suffix.lower()
    if ext not in _ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail="Unsupported file type.")
    contents = await file.read()
    if len(contents) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File must be smaller than 10 MB.")
    string_count = None
    if ext == ".xlsx":
        try:
            from io import BytesIO
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(contents), read_only=True, data_only=True)
            ws = wb.active
            string_count = sum(
                1 for row in ws.iter_rows(min_row=2, max_col=1, values_only=True)
                if row[0] and str(row[0]).strip()
            )
            wb.close()
        except Exception:
            pass
    return JSONResponse({"filename": fname, "file_type": ext.lstrip(".").upper(), "string_count": string_count})


@app.post("/api/translate")
async def start_translation(
    file: UploadFile = File(...),
    tone: str = Form("informal"),
    languages: str = Form("fr,de,pt,ja,es"),
    optimize_length: str = Form("true"),
):
    """Accept .xlsx / .pdf / .docx, convert if needed, kick off translation."""
    fname = file.filename or ""
    ext   = Path(fname).suffix.lower()
    if ext not in _ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail="Supported formats: .xlsx, .pdf, .docx")

    contents = await file.read()
    if len(contents) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File must be smaller than 10 MB.")

    job_id    = uuid.uuid4().hex
    safe_name = f"{job_id}_{fname}"
    raw_path  = Path("uploads") / safe_name
    raw_path.write_bytes(contents)

    loop = asyncio.get_running_loop()

    if ext in (".docx", ".doc"):
        # Native docx pipeline — preserves document structure, outputs a zip
        expected_output = Path("outputs") / f"{raw_path.stem}_translated.zip"
        JOBS[job_id] = {
            "status": "pending",
            "file_type": "docx",
            "input_file": fname,
            "upload_path": str(raw_path),
            "output_path": str(expected_output),
            "created_at": datetime.now().isoformat(),
            "cancel_requested": False,
            "tone": tone,
            "languages": languages,
            "optimize_length": optimize_length.lower() == "true",
            "error": None,
            "review_data": None,
            "report": None,
        }
        loop.run_in_executor(_executor, _run_docx_job, job_id, str(raw_path), languages)
    else:
        # Excel / PDF pipeline — convert if needed, output translated xlsx
        if ext != ".xlsx":
            xlsx_path = raw_path.with_suffix(".xlsx")
            try:
                _convert_to_xlsx(raw_path, xlsx_path)
            except Exception as exc:
                raise HTTPException(status_code=422, detail=f"Could not parse file: {exc}")
            work_path = xlsx_path
        else:
            work_path = raw_path

        expected_output = Path("outputs") / f"{work_path.stem}_translated.xlsx"
        JOBS[job_id] = {
            "status": "pending",
            "file_type": "xlsx",
            "input_file": fname,
            "upload_path": str(work_path),
            "output_path": str(expected_output),
            "created_at": datetime.now().isoformat(),
            "cancel_requested": False,
            "tone": tone,
            "languages": languages,
            "optimize_length": optimize_length.lower() == "true",
            "error": None,
            "review_data": None,
            "report": None,
        }
        loop.run_in_executor(_executor, _run_job, job_id, str(work_path), languages)

    return JSONResponse({"job_id": job_id, "status": "pending"})


@app.get("/api/status/{job_id}")
async def get_status(job_id: str):
    """Return current job status plus review data when complete."""
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")
    total = job.get("total_batches")
    current = job.get("current_batch")
    batch_progress = f"{current}/{total}" if total and total > 1 else None
    return JSONResponse({
        "job_id": job_id,
        "status": job["status"],
        "file_type": job.get("file_type", "xlsx"),
        "input_file": job["input_file"],
        "created_at": job["created_at"],
        "error": job["error"],
        "review_data": job["review_data"],
        "batch_progress": batch_progress,
    })


@app.delete("/api/cancel/{job_id}")
async def cancel_job(job_id: str):
    """Mark a pending/running job as cancelled. The background thread will finish
    on its own but its result will be discarded."""
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job["status"] in ("complete", "failed"):
        raise HTTPException(status_code=400, detail="Job has already finished and cannot be cancelled.")
    job["cancel_requested"] = True
    job["status"] = "cancelled"
    return JSONResponse({"job_id": job_id, "status": "cancelled"})


@app.get("/api/download/{job_id}")
async def download(job_id: str):
    """Stream the translated Excel file to the browser."""
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job["status"] != "complete":
        raise HTTPException(status_code=400, detail="Translation is not complete yet.")

    output_path = Path(job["output_path"])
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Output file not found on disk.")

    input_stem = Path(job["input_file"]).stem
    if job.get("file_type") == "docx":
        return FileResponse(
            path=str(output_path),
            filename=f"{input_stem}_translated.zip",
            media_type="application/zip",
        )
    return FileResponse(
        path=str(output_path),
        filename=f"{input_stem}_translated.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# ---------------------------------------------------------------------------
# Background job runner
# ---------------------------------------------------------------------------

_DOCX_BATCH_SIZE = 10   # segments per direct-LLM translation call for large documents
EXCEL_BATCH_SIZE = 75   # rows per CrewAI kickoff for Excel translation


def _count_excel_rows(path: str) -> int:
    """Count non-empty English-column rows (excluding header) in an Excel file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        ws = wb.active
        count = sum(1 for row in ws.iter_rows(min_row=2, max_col=1, values_only=True) if row[0])
        wb.close()
        return count
    except Exception:
        return 0

_LANG_RULES = {
    "fr":     "French — formal 'vous' register",
    "de":     "German — formal 'Sie' register",
    "pt_BR":  "Brazilian Portuguese — 'você' register",
    "ja":     "Japanese — 丁寧語 polite register",
    "es_ES": "Spain Spanish (Castilian) — 'tú' register throughout UI; 'usted' only in legal/contractual text",
}


def _extract_json_array(text: str) -> str:
    """Extract the first JSON array from text, stripping LLM preamble and markdown fences."""
    text = text.strip()
    # Strip a leading code fence before searching for '['
    text = re.sub(r'^```+(?:json)?\s*', '', text, flags=re.IGNORECASE)
    start = text.find("[")
    if start == -1:
        return text  # let json.loads raise a clear error
    text = text[start:]
    # Strip trailing code fence
    fence_pos = text.rfind("```")
    if fence_pos != -1:
        text = text[:fence_pos].rstrip()
    return text.strip()


def _translate_segment_batch(
    segments: list[dict],
    target_languages: list[str],
    brand_context: str,
) -> list[dict]:
    """Translate one batch of segments via direct Anthropic API call.

    Bypasses CrewAI agents to avoid output-token truncation on large documents.
    """
    import litellm
    import os as _os
    from dotenv import load_dotenv
    load_dotenv()

    model = (
        _os.environ.get("MODEL")
        or _os.environ.get("ANTHROPIC_MODEL")
        or _os.environ.get("OPENAI_MODEL_NAME")
        or "anthropic/claude-sonnet-4-6"
    )

    lang_rules = "\n".join(
        f"  - {_LANG_RULES.get(lc, lc)}" for lc in target_languages
    )
    lang_example = ", ".join(
        f'"{lc}": "<{_LANG_RULES.get(lc, lc).split("—")[0].strip()} translation>"'
        for lc in target_languages
    )
    segs_json = json.dumps(segments, ensure_ascii=False, indent=2)

    prompt = f"""You are translating customer-facing document segments for Monotype.

LANGUAGE REGISTER RULES:
{lang_rules}

KEY BRAND RULES:
- Preserve ALL product names in English: Monotype, MyFonts, Font Manager, Mosaic, Fonts.com, WhatTheFont, Monotype Fonts, etc.
- These are complete sentences from customer communications — translate faithfully, do NOT summarise.
- Placeholder tokens in {{single}} or {{{{double}}}} braces must be copied verbatim.

BRAND CONTEXT SUMMARY:
{brand_context[:2500]}

SEGMENTS TO TRANSLATE (you MUST output ALL {len(segments)} of them):
{segs_json}

TARGET LANGUAGE CODES: {target_languages}

Output ONLY a valid JSON array. Each element MUST follow this exact structure:
{{"segment_id": <int>, "english": "<original text>", {lang_example}}}

Include ONLY the requested language keys. Raw JSON only — no markdown fences, no preamble."""

    response = litellm.completion(
        model=model,
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )

    usage = getattr(response, "usage", None)
    if usage:
        print(
            f"[DocxBatch] tokens — prompt: {usage.prompt_tokens}, "
            f"completion: {usage.completion_tokens}, "
            f"total: {usage.total_tokens}"
        )

    raw = _extract_json_array(response.choices[0].message.content)
    result = json.loads(raw)
    # Attach usage so the caller can accumulate totals
    return result, usage


def _run_docx_job_batched(
    job_id: str,
    docx_path: str,
    segments: list[dict],
    target_languages: list[str],
) -> None:
    """Translate a large docx in batches, then write translated docx files directly."""
    from .crew import DocxTranslationCrew
    from .tools.docx_tools import write_translations_to_docx_impl

    # Step 1: Ensure brand context is cached
    brand_cache = Path("outputs/brand_context_cache.md")
    if brand_cache.exists():
        brand_context = brand_cache.read_text(encoding="utf-8")
    else:
        DocxTranslationCrew()._run_brand_context_only("knowledge")
        brand_context = brand_cache.read_text(encoding="utf-8") if brand_cache.exists() else ""

    # Step 2: Translate in batches
    all_translations: list[dict] = []
    batches = [segments[i:i + _DOCX_BATCH_SIZE] for i in range(0, len(segments), _DOCX_BATCH_SIZE)]
    batch_errors: list[str] = []
    total_prompt_tokens = 0
    total_completion_tokens = 0

    for idx, batch in enumerate(batches, 1):
        print(f"[DocxBatch] Translating batch {idx}/{len(batches)} ({len(batch)} segments)")
        try:
            batch_result, usage = _translate_segment_batch(batch, target_languages, brand_context)
            all_translations.extend(batch_result)
            if usage:
                total_prompt_tokens += usage.prompt_tokens or 0
                total_completion_tokens += usage.completion_tokens or 0
            print(f"[DocxBatch] Batch {idx} OK — {len(batch_result)} entries")
        except Exception as exc:
            err_msg = f"Batch {idx}/{len(batches)} {type(exc).__name__}: {exc}"
            print(f"[DocxBatch] FAILED: {err_msg}")
            batch_errors.append(err_msg)
            JOBS[job_id]["error"] = "; ".join(batch_errors)
            for seg in batch:
                entry: dict = {"segment_id": seg["segment_id"], "english": seg.get("text", "")}
                for lang in target_languages:
                    entry[lang] = seg.get("text", "")
                all_translations.append(entry)

    print(
        f"[DocxBatch] TOTAL tokens — prompt: {total_prompt_tokens}, "
        f"completion: {total_completion_tokens}, "
        f"total: {total_prompt_tokens + total_completion_tokens}"
    )
    JOBS[job_id]["token_usage"] = {
        "prompt_tokens": total_prompt_tokens,
        "completion_tokens": total_completion_tokens,
        "total_tokens": total_prompt_tokens + total_completion_tokens,
    }

    # Step 3: Write merged translations JSON
    review_path = Path("outputs/reviewed_docx_translations.json")
    review_path.write_text(
        json.dumps(all_translations, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"[DocxBatch] Wrote {len(all_translations)} entries to reviewed_docx_translations.json")

    # Step 4: Write translated docx files + zip directly (no AI agent needed)
    result = write_translations_to_docx_impl(docx_path)
    if not result.get("success"):
        raise RuntimeError(f"write_translations_to_docx_impl failed: {result.get('error')}")


def _run_docx_job(job_id: str, docx_path: str, languages: str) -> None:
    """Run the docx translation pipeline in a background thread."""
    JOBS[job_id]["status"] = "running"
    try:
        from .crew import DocxTranslationCrew
        from .tools.docx_tools import extract_segments

        _lang_normalise = {"pt": "pt_BR", "es": "es_ES"}
        target_languages = [
            _lang_normalise.get(l.strip(), l.strip())
            for l in languages.split(",") if l.strip()
        ]

        segments = extract_segments(docx_path)

        if len(segments) > _DOCX_BATCH_SIZE:
            # Large document — batch translation to avoid output-token truncation
            _run_docx_job_batched(job_id, docx_path, segments, target_languages)
        else:
            DocxTranslationCrew().crew().kickoff(inputs={
                "docx_path": docx_path,
                "knowledge_dir": "knowledge",
                "target_languages": target_languages,
            })

        # Locate the zip produced by write_translations_to_docx
        stem = Path(docx_path).stem
        zip_path = Path("outputs") / f"{stem}_translated.zip"
        if zip_path.exists():
            JOBS[job_id]["output_path"] = str(zip_path)

        # Capture review data
        review_path = Path("outputs/reviewed_docx_translations.json")
        if review_path.exists():
            try:
                JOBS[job_id]["review_data"] = json.loads(review_path.read_text())
            except Exception:
                pass

        if not JOBS[job_id].get("cancel_requested"):
            JOBS[job_id]["status"] = "complete"

    except Exception as exc:
        if not JOBS[job_id].get("cancel_requested"):
            JOBS[job_id]["status"] = "failed"
            JOBS[job_id]["error"] = str(exc)


def _run_job(job_id: str, excel_path: str, languages: str = "fr,de,pt,ja,es") -> None:
    """Run the full CrewAI pipeline in a background thread, batching large files."""
    JOBS[job_id]["status"] = "running"
    try:
        from .crew import MonotypeTranslationCrew  # import here to keep startup fast

        _lang_normalise = {"pt": "pt_BR", "es": "es_ES"}
        target_languages = [
            _lang_normalise.get(l.strip(), l.strip())
            for l in languages.split(",") if l.strip()
        ]

        inputs = {
            "excel_path": excel_path,
            "knowledge_dir": "knowledge",
            "target_languages": target_languages,
            "tone": JOBS[job_id].get("tone", "informal"),
        }

        # Compute number of batches so the UI can show "Batch N/M" progress.
        total_rows = _count_excel_rows(excel_path)
        n_batches = max(1, math.ceil(total_rows / EXCEL_BATCH_SIZE))
        JOBS[job_id]["total_batches"] = n_batches

        for batch_num in range(n_batches):
            if JOBS[job_id].get("cancel_requested"):
                break

            JOBS[job_id]["current_batch"] = batch_num + 1
            print(f"[ExcelBatch] Starting batch {batch_num + 1}/{n_batches} for job {job_id}")
            MonotypeTranslationCrew().crew().kickoff(inputs=inputs)

        # Snapshot review data from the last batch (shared file; capture before another job runs)
        review_path = Path("outputs/reviewed_translations.json")
        if review_path.exists():
            try:
                JOBS[job_id]["review_data"] = json.loads(review_path.read_text())
            except Exception:
                pass

        # Capture the latest production report
        reports = sorted(Path("outputs").glob("translation_report-*.md"))
        if reports:
            try:
                JOBS[job_id]["report"] = reports[-1].read_text()
            except Exception:
                pass

        # Only mark complete if the user hasn't cancelled while the crew was running
        if not JOBS[job_id].get("cancel_requested"):
            JOBS[job_id]["status"] = "complete"

    except Exception as exc:
        if not JOBS[job_id].get("cancel_requested"):
            JOBS[job_id]["status"] = "failed"
            JOBS[job_id]["error"] = str(exc)
